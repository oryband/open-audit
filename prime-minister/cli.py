#!/usr/bin/env python

import click
from docx import Document


def build_heirarchy(paragraphs):
    """WORK IN PROGRESS

    walk through all paragraphs and build a paragraph heirarcy.

    font size heuristics are used to determine to which higher-level
    "topic" the current paragraph belongs to.

    e.g. a paragraph with a "normal" (small) font size belongs to the first
    paragraph found before it with a bigger font size.

    returns a dictionary with the highest-level paragraphs (topics),
    whose keys are a list of dictionaries of the same structure,
    but one level below in the heirarchy:

    {
        "topic 1":
            [
                {
                    "body": [ recursive structure ],
                    "topic_name": "topic 1.1",
                },
                {
                    "topic_name": "topic 1.2",
                    "body":
                    ..
                },
                ..
            ],
        "topic 2": [ .. ],
        ..
    }
    """
    pass


@click.group()
def cli(): pass


@cli.command()
@click.argument('document', type=click.Path(exists=True))
@click.option('-v', '--verbose', is_flag=True)
def print_styles(document, verbose):
    for s in paragraphs_by_attr(document,
                                lambda p: set([p.style.name]), verbose):
        click.echo(s)


@cli.command()
@click.argument('document', type=click.Path(exists=True))
@click.option('-v', '--verbose', is_flag=True)
def print_font_sizes_pt(document, verbose):
    """prints all paragraphs with given font size (pt)."""
    def attr_func(p):
        sizes = set(str(run.font.size.pt)
                    for run in p.runs
                    if run.font.size is not None)

        sizes |= set(str(run.style.font.size.pt)
                     for run in p.runs
                     if run.style.font.size is not None)

        if p.style.font.size is not None:
            sizes |= set([str(p.style.font.size.pt)])
        return sizes

    for size in paragraphs_by_attr(document, attr_func, verbose):
        click.echo(size)


def paragraphs_by_attr(document, attr_func, verbose):
    """return an attribute from all paragraphs according to given function.

    the attribute function is used for fetching the attribute from every
    paragraph.
    """
    if verbose:
        click.echo('opening %s' % click.format_filename(document), err=True)

    doc = Document(document)
    attrs = set()

    l = len(doc.paragraphs)
    if verbose:
        click.echo('total of %d paragraphs, starting' % l, err=True)
    for i, p in enumerate(doc.paragraphs, start=1):
        attrs |= attr_func(p)

        if verbose and i % 1000 == 0:
            click.echo('processed %d paragraphs' % i, err=True)

    if verbose:
        click.echo('processing complete, total %d paragraphs' % (i+1),
                   err=True)

    return attrs


@cli.command()
@click.argument('document', type=click.Path(exists=True))
@click.argument('style')
@click.option('-v', '--verbose', is_flag=True)
def print_paragraphs_with_style_name(document, style, verbose):
    """print all paragraphs with given style name."""
    ps = compare_paragraphs_by_attr(document,
                                    style,
                                    attr_func=lambda p: p.style.name,
                                    cmp_func=lambda a, b: a == b,
                                    verbose=verbose)
    for p in ps:
        click.echo(p.text)


@cli.command()
@click.argument('document', type=click.Path(exists=True))
@click.argument('size')
@click.option('-v', '--verbose', is_flag=True)
def print_paragraphs_with_font_size_pt(document, size, verbose):
    def attr_func(p):
        sizes = []

        if p.style.font.size is not None:
            sizes += str(p.style.font.size.pt)

        sizes += ([str(run.font.size.pt)
                   for run in p.runs
                   if run.font.size is not None] +

                  [str(run.style.font.size.pt)
                   for run in p.runs
                   if run.style.font.size is not None])

        return sizes[-1] if sizes else ''

    ps = compare_paragraphs_by_attr(
        document,
        size,
        attr_func=attr_func,
        cmp_func=lambda a, b: a == b,
        verbose=verbose)

    for p in ps:
        click.echo(p.text)


def compare_paragraphs_by_attr(document, attr_value, attr_func, cmp_func,
                               verbose):
    if verbose:
        click.echo('opening %s' % click.format_filename(document), err=True)

    paragraphs = []
    doc = Document(document)
    l = len(doc.paragraphs)
    if verbose:
        click.echo('total of %d paragraphs, starting' % l, err=True)
    for i, p in enumerate(doc.paragraphs, start=1):
        if cmp_func(attr_func(p), attr_value):
            paragraphs.append(p)

        if verbose and i % 1000 == 0:
            click.echo('processed %d paragraphs' % i, err=True)

    return paragraphs


if __name__ == '__main__':
    cli()
